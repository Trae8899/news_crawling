import json
import docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def Json2Word_Titlelink(fileRoute):
    path=os.path.dirname(fileRoute)
    name, _ = os.path.splitext(os.path.basename(fileRoute))

    
    # JSON 파일에서 데이터 로드
    with open(fileRoute, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Word 문서 생성
    doc1 = docx.Document()
    # 각 아이템에 대해 제목, 내용, 그리고 원문 링크를 포함하는 텍스트 상자를 추가

    i=0
    for item in data:
        # if item['Link'] and item['Title_Kr'] and item['Summary_Kr']:
        if item['Link'] and item['Title_Kr']:
            i += 1
            
            title1 = doc1.add_heading("",level=1)
            text2 = doc1.add_paragraph(item['Summary_Kr'])
            text5 = doc1.add_paragraph("///")
            text5.alignment=WD_ALIGN_PARAGRAPH.CENTER
            hyperlink = add_hyperlink(title1, item['Link'], str(i)+". "+item['Title_Kr'], None, True)
            
        else:
            continue

    # Word 문서 저장
    doc1_name = os.path.join(path, f"{name}.docx")
    try:
        doc1.save(doc1_name)
    except:
        return

    return doc1_name

def Json2Word_addlink(fileRoute):
    path=os.path.dirname(fileRoute)
    name, _ = os.path.splitext(os.path.basename(fileRoute))
    # JSON 파일에서 데이터 로드
    with open(fileRoute, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Word 문서 생성
    doc1 = docx.Document()
    # 각 아이템에 대해 제목, 내용, 그리고 원문 링크를 포함하는 텍스트 상자를 추가

    i=0
    for item in data:
        # if item['Link'] and item['Title_Kr'] and item['Summary_Kr']:
        if item['Link'] and item['Title_Kr']:
            i += 1
            
            title1 = doc1.add_heading(str(i)+". "+item['Title_Kr'],level=1)
            text2 = doc1.add_paragraph(item['Summary_Kr'])
            text4 = doc1.add_paragraph(item['Link'])
            text5 = doc1.add_paragraph("///")
            text5.alignment=WD_ALIGN_PARAGRAPH.CENTER
            hyperlink = add_hyperlink(text4, item['Link'], item['Link'], None, True)
            
        else:
            continue

    # Word 문서 저장
    doc1_name = os.path.join(path, f"{name}_link.docx")
    try:
        doc1.save(doc1_name)
    except:
        return

    return doc1_name

def read_word_document(file_path):
    document = docx.Document(file_path)
    content = '\n'.join([paragraph.text for paragraph in document.paragraphs])
    return content

